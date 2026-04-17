"""Word (.docx) file handler using python-docx."""

from __future__ import annotations

import io

from docx import Document

from entity_registry import EntityRegistry


def read_word(file_bytes: bytes) -> tuple[Document, str]:
    """Return (Document object, full_text_for_detection)."""
    doc = Document(io.BytesIO(file_bytes))
    texts: list[str] = []

    for para in doc.paragraphs:
        texts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)

    return doc, "\n".join(texts)


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """Replace entities in a paragraph while preserving run formatting.

    Strategy: join all run texts, do replacements on the joined string,
    then redistribute the result across the original runs.  The last run
    absorbs any length difference.
    """
    if not paragraph.runs:
        return

    full = "".join(run.text for run in paragraph.runs)
    replaced = full
    for orig, pseudo in mapping.items():
        replaced = replaced.replace(orig, pseudo)

    if replaced == full:
        return

    # Redistribute text across runs
    runs = paragraph.runs
    pos = 0
    for i, run in enumerate(runs):
        if i == len(runs) - 1:
            # Last run gets whatever remains
            run.text = replaced[pos:]
        else:
            chunk_len = len(run.text)
            run.text = replaced[pos : pos + chunk_len]
            pos += chunk_len


def write_word(
    doc: Document,
    registry: EntityRegistry,
    skip_entities: set[str] | None = None,
) -> bytes:
    """Replace entities in the document and return docx bytes."""
    if skip_entities is None:
        skip_entities = set()

    mapping = {
        row["original"]: row["pseudonym"]
        for row in registry.mapping_table
        if row["original"] not in skip_entities
    }

    for para in doc.paragraphs:
        _replace_in_paragraph(para, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, mapping)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
