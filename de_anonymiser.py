"""
De-Anonymiser page — reverses anonymisation using a mapping key.

Upload an anonymised file + mapping key (CSV/JSON) → get the original restored.
"""

from __future__ import annotations

import csv
import io
import json
import os

import streamlit as st

from utils import detect_file_type
from file_io.text_handler import read_text, write_text
from file_io.excel_handler import read_excel
from file_io.word_handler import read_word
from file_io.pdf_handler import read_pdf, is_scanned_pdf

import fitz  # PyMuPDF
from openpyxl import Workbook
from docx import Document


# ──────────────────────────────────────────────────────────────────────
# Mapping file parsers
# ──────────────────────────────────────────────────────────────────────
def _parse_mapping_csv(raw: bytes) -> dict[str, str]:
    """Parse mapping CSV → {pseudonym: original}."""
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1")
    reader = csv.DictReader(io.StringIO(text))
    mapping: dict[str, str] = {}
    for row in reader:
        pseudo = row.get("pseudonym", "").strip()
        orig = row.get("original", "").strip()
        if pseudo and orig:
            mapping[pseudo] = orig
    return mapping


def _parse_mapping_json(raw: bytes) -> dict[str, str]:
    """Parse mapping JSON → {pseudonym: original}."""
    data = json.loads(raw)
    mapping: dict[str, str] = {}
    for entry in data:
        pseudo = entry.get("pseudonym", "").strip()
        orig = entry.get("original", "").strip()
        if pseudo and orig:
            mapping[pseudo] = orig
    return mapping


def _parse_mapping(raw: bytes, filename: str) -> dict[str, str]:
    """Auto-detect CSV vs JSON and parse."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".json":
        return _parse_mapping_json(raw)
    return _parse_mapping_csv(raw)


# ──────────────────────────────────────────────────────────────────────
# Generic text replacement (longest-first to avoid partial matches)
# ──────────────────────────────────────────────────────────────────────
def _apply_reverse_mapping(text: str, mapping: dict[str, str]) -> str:
    """Replace pseudonyms with originals, longest pseudonym first."""
    for pseudo, orig in sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True):
        text = text.replace(pseudo, orig)
    return text


# ──────────────────────────────────────────────────────────────────────
# File-type-specific de-anonymisation
# ──────────────────────────────────────────────────────────────────────
def _deanon_text_file(raw: bytes, mapping: dict[str, str]) -> bytes:
    text = read_text(raw)
    return write_text(_apply_reverse_mapping(text, mapping))


def _deanon_csv_file(raw: bytes, mapping: dict[str, str]) -> bytes:
    import pandas as pd
    try:
        df = pd.read_csv(io.BytesIO(raw), dtype=str, keep_default_na=False)
    except UnicodeDecodeError:
        df = pd.read_csv(io.BytesIO(raw), dtype=str, keep_default_na=False, encoding="latin-1")

    sorted_mapping = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)

    def _replace_cell(val: str) -> str:
        for pseudo, orig in sorted_mapping:
            val = val.replace(pseudo, orig)
        return val

    df = df.map(_replace_cell)
    buf = io.BytesIO()
    df.to_csv(buf, index=False, encoding="utf-8")
    return buf.getvalue()


def _deanon_excel_file(raw: bytes, mapping: dict[str, str]) -> bytes:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(raw))
    sorted_mapping = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None and isinstance(cell.value, str):
                    val = cell.value
                    for pseudo, orig in sorted_mapping:
                        val = val.replace(pseudo, orig)
                    cell.value = val

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _deanon_word_file(raw: bytes, mapping: dict[str, str]) -> bytes:
    doc = Document(io.BytesIO(raw))
    sorted_mapping = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)

    def _replace_in_paragraph(para):
        if not para.runs:
            return
        full = "".join(run.text for run in para.runs)
        replaced = full
        for pseudo, orig in sorted_mapping:
            replaced = replaced.replace(pseudo, orig)
        if replaced == full:
            return
        runs = para.runs
        pos = 0
        for i, run in enumerate(runs):
            if i == len(runs) - 1:
                run.text = replaced[pos:]
            else:
                chunk_len = len(run.text)
                run.text = replaced[pos : pos + chunk_len]
                pos += chunk_len

    for para in doc.paragraphs:
        _replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _deanon_pdf_file(raw: bytes, mapping: dict[str, str]) -> bytes:
    doc = fitz.open(stream=raw, filetype="pdf")
    sorted_mapping = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)

    for page in doc:
        for pseudo, orig in sorted_mapping:
            rects = page.search_for(pseudo)
            for rect in rects:
                page.add_redact_annot(rect, text=orig, fontsize=0)
        page.apply_redactions()

    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────
# Streamlit page
# ──────────────────────────────────────────────────────────────────────
def render():
    """Render the De-Anonymise page."""
    st.header("De-Anonymise")
    st.caption("Upload an anonymised file and its mapping key to restore original content.")

    col1, col2 = st.columns(2)

    with col1:
        input_file = st.file_uploader(
            "Anonymised file",
            type=["pdf", "xlsx", "docx", "csv", "txt"],
            key="deanon_input",
        )

    with col2:
        mapping_file = st.file_uploader(
            "Mapping key (CSV or JSON)",
            type=["csv", "json"],
            key="deanon_mapping",
        )

    if input_file and mapping_file:
        # Parse mapping
        mapping_raw = mapping_file.read()
        try:
            mapping = _parse_mapping(mapping_raw, mapping_file.name)
        except Exception as e:
            st.error(f"Failed to parse mapping file: {e}")
            return

        if not mapping:
            st.warning("Mapping file is empty or has no valid entries.")
            return

        # Show mapping preview
        import pandas as pd
        with st.expander(f"Mapping preview ({len(mapping)} entries)", expanded=False):
            preview_df = pd.DataFrame(
                [{"pseudonym": k, "original": v} for k, v in mapping.items()]
            )
            st.dataframe(preview_df, use_container_width=True)

        # De-anonymise
        if st.button("De-Anonymise", type="primary"):
            input_raw = input_file.read()
            ftype = detect_file_type(input_file.name)

            with st.spinner("Restoring original content..."):
                try:
                    if ftype == "text":
                        out_bytes = _deanon_text_file(input_raw, mapping)
                    elif ftype == "csv":
                        out_bytes = _deanon_csv_file(input_raw, mapping)
                    elif ftype == "excel":
                        out_bytes = _deanon_excel_file(input_raw, mapping)
                    elif ftype == "word":
                        out_bytes = _deanon_word_file(input_raw, mapping)
                    elif ftype == "pdf":
                        out_bytes = _deanon_pdf_file(input_raw, mapping)
                    else:
                        out_bytes = _deanon_text_file(input_raw, mapping)
                except Exception as e:
                    st.error(f"De-anonymisation failed: {e}")
                    return

            st.success("De-anonymisation complete.")

            # Download
            mime_map = {
                "pdf": "application/pdf",
                "excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "csv": "text/csv",
                "text": "text/plain",
            }
            ext_map = {
                "pdf": ".pdf",
                "excel": ".xlsx",
                "word": ".docx",
                "csv": ".csv",
                "text": ".txt",
            }
            base, _ = os.path.splitext(input_file.name)
            # Remove _anonymised suffix if present
            if base.endswith("_anonymised"):
                base = base[: -len("_anonymised")]
            restored_name = f"{base}_restored{ext_map.get(ftype, '.txt')}"

            st.download_button(
                label=f"Download {restored_name}",
                data=out_bytes,
                file_name=restored_name,
                mime=mime_map.get(ftype, "application/octet-stream"),
                key="dl_restored",
            )
